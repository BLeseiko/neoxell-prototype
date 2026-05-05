from docx import Document
import sys

def extract_text(path):
    try:
        doc = Document(path)
        print("--- Paragraphs ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        print("\n--- Tables ---")
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_data))
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_text("/Users/antoninatrofimchuk/Desktop/Neoxell/neoxell-prototype/NeoXell_Website_ Concept_Feedback.docx")
